from pathlib import Path

import docx
import pdfplumber


def extract_text(file_path):
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        text = _extract_pdf(path)
    elif suffix == ".docx":
        text = _extract_docx(path)
    elif suffix == ".txt":
        text = path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported resume format: {suffix} (supported: .pdf, .docx, .txt)")

    # Some PDF fonts encode dashes in a way pdfplumber can't decode; normalize
    # the resulting replacement character to a plain hyphen.
    return text.replace("�", "-")


def _extract_pdf(path):
    chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                chunks.append(text)
    return "\n".join(chunks)


def _extract_docx(path):
    document = docx.Document(path)
    return "\n".join(p.text for p in document.paragraphs if p.text)


def extract_links(file_path):
    """Pull hyperlink URLs that are embedded as clickable links rather than
    plain visible text (e.g. "Live: Link" anchors) - these are invisible to
    extract_text() and to any plain-text URL regex."""
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        links = _extract_pdf_links(path)
    elif suffix == ".docx":
        links = _extract_docx_links(path)
    else:
        links = []

    # Hyperlink annotations can be mailto:, tel:, etc. - not verifiable project links.
    return [link for link in links if link.startswith(("http://", "https://"))]


def _extract_pdf_links(path):
    links = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            for link in page.hyperlinks:
                uri = link.get("uri")
                if uri:
                    links.append(uri)
    return links


def _extract_docx_links(path):
    links = []
    try:
        document = docx.Document(path)
        for rel in document.part.rels.values():
            if "hyperlink" in rel.reltype and rel.target_ref.startswith("http"):
                links.append(rel.target_ref)
    except Exception:
        pass
    return links
